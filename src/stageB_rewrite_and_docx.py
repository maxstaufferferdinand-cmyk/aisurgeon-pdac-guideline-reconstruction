#!/usr/bin/env python3
"""
FINAL STAGE B — Chapter rewrite + Word DOCX assembly.

Purpose
-------
Create a text-first living-evidence update that follows the chronology of the
2015 ESMO pancreatic-cancer guideline.

VERSION 1 CONTENT DESIGN
------------------------
For each original chapter, in original chronological order:

1. ORIGINAL ESMO 2015
   - inserted deterministically from esmo2015_baseline_by_chapter.json
   - original citation numbers [1]-[58] remain untouched
   - GPT is NOT allowed to rewrite or paraphrase this original block

2. EVIDENCE UPDATE 2015-31 AUG 2023
   - generated from the completed Stage-A evidence-unit syntheses
   - new claims cite new PubMed evidence only
   - citations are returned as PMIDs and numbered deterministically from [59]
     onward during DOCX assembly

3. CURRENT CLINICAL PRACTICE
   - concise clinically oriented summary
   - citations restricted to Stage-A MAIN_SYNTHESIS papers
   - evidence can confirm the 2015 position or support updated practice

4. APPENDIX
   - default: included at the end of the same DOCX
   - summarizes APPENDIX/context evidence that should not drive clinical practice
   - rejected evidence remains in the Stage-A audit files and is not promoted
     into the main narrative

Figures/flowcharts are intentionally omitted in v1 and can be regenerated later.

Inputs
------
data/stageA_evidence_synthesis_manifest.json
data/stageA_unit_evidence_synthesis.jsonl
data/guideline_integration_master_v2.jsonl
data/final_evidence_assignments_v2.jsonl
data/esmo2015_baseline_by_chapter.json

Outputs
-------
data/stageB_chapter_rewrite_batch_input.jsonl
data/stageB_chapter_rewrite_batch_output.jsonl
data/stageB_chapter_rewrite_results.jsonl
data/stageB_chapter_rewrite_parse_failures.jsonl
data/stageB_chapter_rewrite_manifest.json
data/stageB_chapter_rewrite_state.json
data/stageB_reference_registry.json

Word:
output/ESMO_PDAC_2015_Living_Evidence_Update_2023_v1.docx

Optional separate appendix:
output/ESMO_PDAC_2015_Living_Evidence_Update_2023_v1_APPENDIX.docx

Model default:
    gpt-5.6-sol
Reasoning:
    high
"""

from __future__ import annotations

import argparse
import json
import os
import re
import time
from collections import Counter, defaultdict
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterable

import requests

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
OUTPUT = ROOT / "output"

STAGEA_MANIFEST = DATA / "stageA_evidence_synthesis_manifest.json"
STAGEA_SYNTHESIS = DATA / "stageA_unit_evidence_synthesis.jsonl"
INTEGRATION_MASTER = DATA / "guideline_integration_master_v2.jsonl"
FINAL_ASSIGNMENTS = DATA / "final_evidence_assignments_v2.jsonl"
BASELINE = DATA / "esmo2015_baseline_by_chapter.json"

BATCH_INPUT = DATA / "stageB_chapter_rewrite_batch_input.jsonl"
BATCH_OUTPUT = DATA / "stageB_chapter_rewrite_batch_output.jsonl"
BATCH_ERRORS = DATA / "stageB_chapter_rewrite_batch_errors.jsonl"
STATE = DATA / "stageB_chapter_rewrite_state.json"

RESULTS = DATA / "stageB_chapter_rewrite_results.jsonl"
FAILURES = DATA / "stageB_chapter_rewrite_parse_failures.jsonl"
MANIFEST = DATA / "stageB_chapter_rewrite_manifest.json"
REFERENCE_REGISTRY = DATA / "stageB_reference_registry.json"

DOCX_PATH = OUTPUT / "ESMO_PDAC_2015_Living_Evidence_Update_2023_v1.docx"
APPENDIX_DOCX_PATH = OUTPUT / "ESMO_PDAC_2015_Living_Evidence_Update_2023_v1_APPENDIX.docx"

OPENAI_BASE_URL = "https://api.openai.com/v1"
TERMINAL = {"completed", "failed", "expired", "cancelled"}
TRANSIENT = {408, 409, 429, 500, 502, 503, 504}

CHAPTER_ORDER = ["1", "2", "3", "4.1", "4.2", "4.3", "5", "6"]

CHAPTER_FALLBACK_TITLES = {
    "1": "Incidence and epidemiology",
    "2": "Diagnosis and pathology/molecular biology",
    "3": "Staging and risk assessment",
    "4.1": "Treatment of localised disease",
    "4.2": "Treatment of non-resectable disease - borderline resectable / locally advanced",
    "4.3": "Treatment of advanced/metastatic disease",
    "5": "Personalised medicine",
    "6": "Follow-up and long-term implications",
}

STAGEB_SYSTEM_PROMPT = r"""
You are writing the evidence-update layer for a living clinical-practice
guideline on pancreatic cancer.

SOURCE FRAME
The underlying backbone is the 2015 ESMO Clinical Practice Guideline.
The original 2015 text will be inserted into the Word document separately and
deterministically. DO NOT rewrite, paraphrase, reproduce, modernize, or silently
correct the original ESMO text. Your output is ONLY the NEW evidence-update
layer and current clinical-practice interpretation.

The evidence search covers 2015 through 31 August 2023.

INPUT EVIDENCE
You receive final Stage-A evidence-unit syntheses. These have already appraised
mapped papers and classified them as:
- MAIN_SYNTHESIS
- CONTEXT_ONLY
- APPENDIX
- REJECT

The Stage-A synthesis already applied the user-defined hierarchy:
Tier 1: meta-analysis of HUMAN RCTs
Tier 2: meta-analysis of HUMAN retrospective/non-randomized or mixed studies
Tier 3: systematic review
Tier 4: other review OR standalone RCT

Do not redo the entire screening from scratch. Respect the Stage-A appraisal,
its uncertainty, and its evidence hierarchy.

MAIN WRITING RULES
1. Follow the chronological/clinical order of the supplied original 2015
   chapter.
2. Write in polished English suitable for a clinical-practice guideline.
3. Clearly distinguish:
   a) what newer evidence confirms,
   b) what it modifies,
   c) what is genuinely new,
   d) where evidence remains insufficient or non-translatable.
4. Do not make a clinical claim solely from APPENDIX or REJECT evidence.
5. OTHER_REVIEW evidence may contextualize the narrative, but must not by
   itself justify ADD/MODIFY/REMOVE of clinical practice.
6. For therapeutic/interventional practice, recommendation-driving statements
   require clinically meaningful human evidence; do not elevate surrogate-only
   findings into clinical practice.
7. For epidemiology, diagnosis, staging/prognosis, personalised medicine and
   follow-up, use domain-appropriate clinical endpoint logic.
8. Do not invent effect sizes, trial names, endpoints, populations, or
   recommendations absent from the supplied evidence memos.
9. Do not cite papers outside the supplied Stage-A PMID sets.
10. Do not use original ESMO reference numbers for new evidence. New citations
    are expressed ONLY as PMIDs in the structured output. The document builder
    will assign [59], [60], ... deterministically.
11. Keep the update readable. Do not turn the chapter into an exhaustive
    catalogue of every paper.
12. Preserve important conflicting evidence and limitations.

CITATION RULES
- evidence_update paragraphs may cite MAIN_SYNTHESIS and CONTEXT_ONLY PMIDs.
- current_clinical_practice points may cite MAIN_SYNTHESIS PMIDs ONLY.
- appendix topics may cite APPENDIX or CONTEXT_ONLY PMIDs.
- REJECT PMIDs must not be cited in the rewritten clinical text.

CHANGE SIGNALS
Use these only as internal structured labels:
CONFIRM
MODIFY
ADD
REMOVE
INSUFFICIENT_EVIDENCE

CURRENT CLINICAL PRACTICE
This section is the concise practical synthesis as of evidence through
31 August 2023. If the newer evidence does not justify a change, explicitly say
that the 2015 position remains supported or that no evidence-based practice
change can be made.

Do not create figures, algorithms or flowcharts. Those are intentionally deferred
to a later run.
"""


def clean(v: Any) -> str:
    return " ".join(str(v or "").replace("\ufeff", "").split())


def load_json(path: Path) -> Any:
    if not path.exists():
        raise FileNotFoundError(f"Missing required file: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def load_jsonl(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        raise FileNotFoundError(f"Missing required file: {path}")
    rows = []
    with path.open("r", encoding="utf-8") as f:
        for n, line in enumerate(f, 1):
            if not line.strip():
                continue
            try:
                rows.append(json.loads(line))
            except json.JSONDecodeError as e:
                raise RuntimeError(f"{path.name} line {n}: {e}") from e
    return rows


def write_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def baseline_by_chapter(raw: Any) -> dict[str, Any]:
    if isinstance(raw, dict):
        if all(cid in raw for cid in CHAPTER_ORDER):
            return {cid: raw[cid] for cid in CHAPTER_ORDER}

        chapters = raw.get("chapters")
        if isinstance(chapters, dict) and all(cid in chapters for cid in CHAPTER_ORDER):
            return {cid: chapters[cid] for cid in CHAPTER_ORDER}

        if isinstance(chapters, list):
            mapped = {
                clean(x.get("chapter_id") or x.get("id")): x
                for x in chapters
                if isinstance(x, dict)
            }
            if all(cid in mapped for cid in CHAPTER_ORDER):
                return {cid: mapped[cid] for cid in CHAPTER_ORDER}

    raise RuntimeError(
        "Could not recognize esmo2015_baseline_by_chapter.json chapter structure."
    )


def validate_stageA() -> None:
    m = load_json(STAGEA_MANIFEST)
    status = clean(m.get("status"))
    successful = int(m.get("successful_unit_syntheses", -1))
    failed = int(m.get("failed_or_missing_unit_syntheses", -1))

    if (
        status != "READY_FOR_STAGE_B_GUIDELINE_UPDATE"
        or successful != 184
        or failed != 0
    ):
        raise RuntimeError(
            "Stage A is not complete. Required:\n"
            "  status=READY_FOR_STAGE_B_GUIDELINE_UPDATE\n"
            "  successful_unit_syntheses=184\n"
            "  failed_or_missing_unit_syntheses=0\n"
            f"Found: status={status!r}, successful={successful}, failed={failed}\n"
            "Run the repaired Stage-A merge before Stage B."
        )

    rows = load_jsonl(STAGEA_SYNTHESIS)
    if len(rows) != 184:
        raise RuntimeError(
            f"Expected 184 Stage-A unit syntheses, found {len(rows)}."
        )


def chapter_sets(unit_rows: list[dict[str, Any]]) -> dict[str, dict[str, set[str]]]:
    out = {
        cid: {
            "main": set(),
            "context": set(),
            "appendix": set(),
            "reject": set(),
        }
        for cid in CHAPTER_ORDER
    }

    for row in unit_rows:
        cid = clean(row.get("chapter_id"))
        if cid not in out:
            raise RuntimeError(f"Unexpected Stage-A chapter {cid!r}")
        result = row.get("result") or {}

        for p in result.get("main_synthesis_pmids") or []:
            out[cid]["main"].add(clean(p))
        for p in result.get("context_only_pmids") or []:
            out[cid]["context"].add(clean(p))
        for p in result.get("appendix_pmids") or []:
            out[cid]["appendix"].add(clean(p))
        for p in result.get("rejected_pmids") or []:
            out[cid]["reject"].add(clean(p))

    return out


def compact_unit_synthesis(row: dict[str, Any]) -> dict[str, Any]:
    result = row.get("result") or {}
    return {
        "evidence_unit_id": clean(row.get("evidence_unit_id")),
        "evidence_unit_name": clean(row.get("evidence_unit_name")),
        "evidence_unit_definition": clean(row.get("evidence_unit_definition")),
        "evidence_unit_origin": clean(row.get("evidence_unit_origin")),
        "mapped_evidence_count": row.get("mapped_evidence_count"),
        "main_synthesis_pmids": result.get("main_synthesis_pmids") or [],
        "context_only_pmids": result.get("context_only_pmids") or [],
        "appendix_pmids": result.get("appendix_pmids") or [],
        "rejected_pmids": result.get("rejected_pmids") or [],
        "tier_counts_main_synthesis": result.get("tier_counts_main_synthesis") or {},
        "key_clinical_findings": result.get("key_clinical_findings") or [],
        "evidence_conflicts": result.get("evidence_conflicts") or [],
        "recommendation_readiness": clean(result.get("recommendation_readiness")),
        "potential_guideline_implication": clean(
            result.get("potential_guideline_implication")
        ),
        "recommendation_driving_evidence_summary": clean(
            result.get("recommendation_driving_evidence_summary")
        ),
        "context_evidence_summary": clean(result.get("context_evidence_summary")),
        "appendix_evidence_summary": clean(result.get("appendix_evidence_summary")),
        "clinical_translation_summary": clean(
            result.get("clinical_translation_summary")
        ),
        "limitations": clean(result.get("limitations")),
        "evidence_memo": clean(result.get("evidence_memo")),
    }


def response_schema() -> dict[str, Any]:
    paragraph = {
        "type": "object",
        "properties": {
            "text": {"type": "string"},
            "citation_pmids": {
                "type": "array",
                "items": {"type": "string"},
                "maxItems": 20,
            },
        },
        "required": ["text", "citation_pmids"],
        "additionalProperties": False,
    }

    practice_point = {
        "type": "object",
        "properties": {
            "text": {"type": "string"},
            "citation_pmids": {
                "type": "array",
                "items": {"type": "string"},
                "maxItems": 15,
            },
            "practice_status": {
                "type": "string",
                "enum": [
                    "CURRENT_STANDARD",
                    "REASONABLE_OPTION",
                    "SELECTED_PATIENTS",
                    "NOT_RECOMMENDED",
                    "INSUFFICIENT_EVIDENCE",
                    "RESEARCH_ONLY",
                ],
            },
        },
        "required": ["text", "citation_pmids", "practice_status"],
        "additionalProperties": False,
    }

    return {
        "name": "pdac_stageB_chapter_update",
        "strict": True,
        "schema": {
            "type": "object",
            "properties": {
                "chapter_id": {"type": "string"},
                "chapter_update_title": {"type": "string"},
                "overall_update_summary": {"type": "string"},
                "updated_sections": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "section_heading": {"type": "string"},
                            "anchor_in_original": {"type": "string"},
                            "change_signal": {
                                "type": "string",
                                "enum": [
                                    "CONFIRM",
                                    "MODIFY",
                                    "ADD",
                                    "REMOVE",
                                    "INSUFFICIENT_EVIDENCE",
                                ],
                            },
                            "source_unit_ids": {
                                "type": "array",
                                "items": {"type": "string"},
                                "maxItems": 30,
                            },
                            "update_paragraphs": {
                                "type": "array",
                                "items": paragraph,
                                "maxItems": 20,
                            },
                            "current_clinical_practice": {
                                "type": "array",
                                "items": practice_point,
                                "maxItems": 12,
                            },
                            "interpretation_of_2015_position": {"type": "string"},
                        },
                        "required": [
                            "section_heading",
                            "anchor_in_original",
                            "change_signal",
                            "source_unit_ids",
                            "update_paragraphs",
                            "current_clinical_practice",
                            "interpretation_of_2015_position",
                        ],
                        "additionalProperties": False,
                    },
                    "maxItems": 30,
                },
                "chapter_current_clinical_practice_summary": {
                    "type": "array",
                    "items": practice_point,
                    "maxItems": 20,
                },
                "appendix_topics": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string"},
                            "summary": {"type": "string"},
                            "citation_pmids": {
                                "type": "array",
                                "items": {"type": "string"},
                                "maxItems": 25,
                            },
                        },
                        "required": ["title", "summary", "citation_pmids"],
                        "additionalProperties": False,
                    },
                    "maxItems": 30,
                },
                "chapter_limitations": {"type": "string"},
            },
            "required": [
                "chapter_id",
                "chapter_update_title",
                "overall_update_summary",
                "updated_sections",
                "chapter_current_clinical_practice_summary",
                "appendix_topics",
                "chapter_limitations",
            ],
            "additionalProperties": False,
        },
    }


def prepare(model: str, effort: str) -> None:
    validate_stageA()

    baseline = baseline_by_chapter(load_json(BASELINE))
    units = load_jsonl(STAGEA_SYNTHESIS)
    sets = chapter_sets(units)

    by_chapter = defaultdict(list)
    for row in units:
        by_chapter[clean(row.get("chapter_id"))].append(row)

    requests_out = []

    for cid in CHAPTER_ORDER:
        chapter_units = sorted(
            by_chapter[cid],
            key=lambda x: clean(x.get("evidence_unit_id")),
        )

        user_prompt = "\n".join([
            f"CHAPTER ID: {cid}",
            f"CHAPTER TITLE: {CHAPTER_FALLBACK_TITLES[cid]}",
            "",
            "ORIGINAL ESMO-2015 CHAPTER BASELINE OBJECT:",
            json.dumps(baseline[cid], ensure_ascii=False, indent=2),
            "",
            "FINAL STAGE-A EVIDENCE-UNIT SYNTHeses FOR THIS CHAPTER:",
            json.dumps(
                [compact_unit_synthesis(x) for x in chapter_units],
                ensure_ascii=False,
                indent=2,
            ),
            "",
            "CHAPTER PMID SETS:",
            json.dumps(
                {
                    "MAIN_SYNTHESIS": sorted(sets[cid]["main"]),
                    "CONTEXT_ONLY": sorted(sets[cid]["context"]),
                    "APPENDIX": sorted(sets[cid]["appendix"]),
                    "REJECT": sorted(sets[cid]["reject"]),
                },
                ensure_ascii=False,
                indent=2,
            ),
            "",
            "Write only the NEW evidence-update layer. Follow the original "
            "chapter chronology. Do not reproduce the original 2015 text. "
            "Do not create figures.",
        ])

        body = {
            "model": model,
            "messages": [
                {"role": "system", "content": STAGEB_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            "response_format": {
                "type": "json_schema",
                "json_schema": response_schema(),
            },
            "max_completion_tokens": 50000,
            "reasoning_effort": effort,
        }

        requests_out.append({
            "custom_id": f"stageB-chapter-{cid.replace('.', '_DOT_')}",
            "method": "POST",
            "url": "/v1/chat/completions",
            "body": body,
        })

    write_jsonl(BATCH_INPUT, requests_out)

    size_mb = BATCH_INPUT.stat().st_size / 1024 / 1024
    if size_mb > 190:
        raise RuntimeError(f"Stage-B Batch input too large: {size_mb:.2f} MB")

    print("\nFINAL STAGE B prepared.")
    print(f"  chapter requests:       {len(requests_out)}")
    print(f"  Batch JSONL:            {size_mb:.2f} MB")
    print(f"  model:                  {model}")
    print(f"  reasoning effort:       {effort}")
    print(f"  input:                  {BATCH_INPUT}")


class Client:
    def __init__(self, key: str, retry_wait: int):
        if not key:
            raise RuntimeError("OPENAI_API_KEY is not set.")
        self.s = requests.Session()
        self.headers = {"Authorization": f"Bearer {key}"}
        self.retry_wait = retry_wait

    def request(self, method: str, path: str, *, timeout: int = 1200, **kwargs):
        url = path if path.startswith("http") else OPENAI_BASE_URL + path

        while True:
            headers = dict(self.headers)
            headers.update(kwargs.pop("headers", {}))
            try:
                r = self.s.request(
                    method, url, headers=headers, timeout=timeout, **kwargs
                )
            except (
                requests.Timeout,
                requests.ConnectionError,
                requests.ChunkedEncodingError,
                requests.ContentDecodingError,
            ) as e:
                print(
                    f"WARN {type(e).__name__}: {e}; retry in {self.retry_wait}s"
                )
                time.sleep(self.retry_wait)
                continue

            if r.status_code in TRANSIENT:
                print(
                    f"WARN HTTP {r.status_code}; retry in {self.retry_wait}s"
                )
                time.sleep(self.retry_wait)
                continue

            if r.status_code >= 400:
                raise RuntimeError(
                    f"OpenAI HTTP {r.status_code}: {r.text[:5000]}"
                )
            return r


def load_state() -> dict[str, Any]:
    if not STATE.exists():
        return {}
    return load_json(STATE)


def save_state(state: dict[str, Any]) -> None:
    STATE.write_text(
        json.dumps(state, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def submit(client: Client, model: str, effort: str) -> None:
    state = load_state()
    if state.get("batch_id"):
        print(f"Existing Stage-B Batch: {state['batch_id']}")
        print("Resuming; no duplicate Batch submitted.")
        return

    with BATCH_INPUT.open("rb") as f:
        upload = client.request(
            "POST",
            "/files",
            files={"file": (BATCH_INPUT.name, f, "application/jsonl")},
            data={"purpose": "batch"},
        ).json()

    payload = {
        "input_file_id": upload["id"],
        "endpoint": "/v1/chat/completions",
        "completion_window": "24h",
        "metadata": {
            "project": "ESMO_PDAC_2015_to_2023_PoC",
            "task": "stageB_chapter_rewrite",
            "model": model,
            "reasoning_effort": effort,
        },
    }

    batch = client.request(
        "POST",
        "/batches",
        headers={"Content-Type": "application/json"},
        data=json.dumps(payload),
    ).json()

    state = {
        "input_file_id": upload["id"],
        "batch_id": batch["id"],
        "status": batch.get("status"),
        "model": model,
        "reasoning_effort": effort,
    }
    save_state(state)
    print(f"Batch id: {batch['id']}")
    print(f"Status:   {batch.get('status')}")


def download(client: Client, file_id: str, destination: Path) -> None:
    destination.write_bytes(
        client.request(
            "GET",
            f"/files/{file_id}/content",
            timeout=1200,
        ).content
    )


def watch(client: Client, poll_seconds: int) -> dict[str, Any]:
    state = load_state()
    batch_id = state.get("batch_id")
    if not batch_id:
        raise RuntimeError("No Stage-B batch_id in state.")

    while True:
        batch = client.request("GET", f"/batches/{batch_id}").json()
        status = batch.get("status")
        counts = batch.get("request_counts") or {}

        print(
            f"status={status}; total={counts.get('total')}; "
            f"completed={counts.get('completed')}; "
            f"failed={counts.get('failed')}"
        )

        state.update({
            "status": status,
            "output_file_id": batch.get("output_file_id"),
            "error_file_id": batch.get("error_file_id"),
            "request_counts": counts,
        })
        save_state(state)

        if status in TERMINAL:
            if batch.get("output_file_id"):
                download(client, batch["output_file_id"], BATCH_OUTPUT)
            if batch.get("error_file_id"):
                download(client, batch["error_file_id"], BATCH_ERRORS)
            return batch

        time.sleep(poll_seconds)


def parse_chapter_custom_id(cid: str) -> str:
    prefix = "stageB-chapter-"
    if not cid.startswith(prefix):
        raise ValueError(f"Invalid Stage-B custom_id: {cid}")
    return cid[len(prefix):].replace("_DOT_", ".")


def extract_content(obj: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    if obj.get("error"):
        raise ValueError(f"Batch error: {obj['error']}")
    response = obj.get("response")
    if not response or response.get("status_code") != 200:
        raise ValueError(
            f"HTTP {response.get('status_code') if response else 'missing'}"
        )

    body = response.get("body") or {}
    choice = (body.get("choices") or [{}])[0]
    if clean(choice.get("finish_reason")) != "stop":
        raise ValueError(
            f"finish_reason={choice.get('finish_reason')}"
        )
    content = ((choice.get("message") or {}).get("content") or "")
    if not content.strip():
        raise ValueError("Empty Stage-B content")
    return content, body.get("usage") or {}


def citation_pmids_from_result(parsed: dict[str, Any]) -> dict[str, list[str]]:
    groups = {
        "update": [],
        "practice": [],
        "appendix": [],
    }

    for section in parsed.get("updated_sections") or []:
        for p in section.get("update_paragraphs") or []:
            groups["update"].extend(clean(x) for x in p.get("citation_pmids") or [])
        for p in section.get("current_clinical_practice") or []:
            groups["practice"].extend(clean(x) for x in p.get("citation_pmids") or [])

    for p in parsed.get("chapter_current_clinical_practice_summary") or []:
        groups["practice"].extend(clean(x) for x in p.get("citation_pmids") or [])

    for topic in parsed.get("appendix_topics") or []:
        groups["appendix"].extend(
            clean(x) for x in topic.get("citation_pmids") or []
        )

    return groups


def merge() -> dict[str, Any]:
    validate_stageA()
    units = load_jsonl(STAGEA_SYNTHESIS)
    sets = chapter_sets(units)

    results = []
    failures = []
    seen = set()

    with BATCH_OUTPUT.open("r", encoding="utf-8") as f:
        for line_no, line in enumerate(f, 1):
            if not line.strip():
                continue
            obj = json.loads(line)
            custom_id = clean(obj.get("custom_id"))

            try:
                cid = parse_chapter_custom_id(custom_id)
                if cid not in CHAPTER_ORDER:
                    raise ValueError(f"Unexpected chapter {cid}")
                if cid in seen:
                    raise ValueError(f"Duplicate chapter output {cid}")
                seen.add(cid)

                content, usage = extract_content(obj)
                parsed = json.loads(content)

                if clean(parsed.get("chapter_id")) != cid:
                    raise ValueError(
                        f"Returned chapter_id={parsed.get('chapter_id')} != {cid}"
                    )

                citations = citation_pmids_from_result(parsed)

                update_allowed = sets[cid]["main"] | sets[cid]["context"]
                practice_allowed = sets[cid]["main"]
                # Appendix may cite any Stage-A evidence that was retained as
                # MAIN_SYNTHESIS, CONTEXT_ONLY, or APPENDIX. REJECT remains forbidden.
                # This allows clinically important MAIN_SYNTHESIS papers to be
                # discussed again in the appendix without downgrading them.
                appendix_allowed = (
                    sets[cid]["main"]
                    | sets[cid]["context"]
                    | sets[cid]["appendix"]
                )
                rejected = sets[cid]["reject"]

                bad_update = set(citations["update"]) - update_allowed
                bad_practice = set(citations["practice"]) - practice_allowed
                bad_appendix = set(citations["appendix"]) - appendix_allowed
                rejected_cited = (
                    set(citations["update"])
                    | set(citations["practice"])
                    | set(citations["appendix"])
                ) & rejected

                if bad_update:
                    raise ValueError(
                        f"Update cites disallowed PMIDs: {sorted(bad_update)}"
                    )
                if bad_practice:
                    raise ValueError(
                        f"Clinical-practice section cites non-MAIN PMIDs: "
                        f"{sorted(bad_practice)}"
                    )
                if bad_appendix:
                    raise ValueError(
                        f"Appendix cites disallowed PMIDs: {sorted(bad_appendix)}"
                    )
                if rejected_cited:
                    raise ValueError(
                        f"Rejected PMIDs cited: {sorted(rejected_cited)}"
                    )

                results.append({
                    "chapter_id": cid,
                    "chapter_title": CHAPTER_FALLBACK_TITLES[cid],
                    "model": load_state().get("model", "gpt-5.6-sol"),
                    "reasoning_effort": load_state().get(
                        "reasoning_effort", "high"
                    ),
                    "usage": usage,
                    "result": parsed,
                })

            except Exception as e:
                failures.append({
                    "line": line_no,
                    "custom_id": custom_id,
                    "error": f"{type(e).__name__}: {e}",
                })

    for cid in CHAPTER_ORDER:
        if cid not in seen:
            failures.append({
                "line": None,
                "custom_id": f"stageB-chapter-{cid.replace('.', '_DOT_')}",
                "error": "Missing chapter output",
            })

    results.sort(key=lambda x: CHAPTER_ORDER.index(x["chapter_id"]))
    write_jsonl(RESULTS, results)
    write_jsonl(FAILURES, failures)

    manifest = {
        "status": (
            "READY_FOR_DOCX"
            if len(results) == 8 and not failures
            else "INCOMPLETE"
        ),
        "chapter_results": len(results),
        "parse_or_validation_failures": len(failures),
        "model": load_state().get("model", "gpt-5.6-sol"),
        "reasoning_effort": load_state().get("reasoning_effort", "high"),
        "document_design": (
            "Original ESMO 2015 content remains deterministic and unchanged; "
            "Stage B supplies only the evidence-update and current-practice layer."
        ),
        "figures": "DEFERRED_TO_SECOND_RUN",
        "new_reference_numbering": "Starts at [59] and follows first appearance.",
    }
    MANIFEST.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print("\nFINAL STAGE B merge completed.")
    print(f"  successful chapters:       {len(results)}/8")
    print(f"  failures:                  {len(failures)}")
    print(f"  results:                   {RESULTS}")
    print(f"  failures file:             {FAILURES}")
    print(f"  manifest:                  {MANIFEST}")

    if failures:
        print("\nWARNING: repair Stage-B failures before DOCX generation.")

    return manifest


def walk_objects(x: Any) -> Iterable[Any]:
    yield x
    if isinstance(x, dict):
        for v in x.values():
            yield from walk_objects(v)
    elif isinstance(x, list):
        for v in x:
            yield from walk_objects(v)


def walk_key_values(x: Any) -> Iterable[tuple[str, Any]]:
    if isinstance(x, dict):
        for k, v in x.items():
            yield str(k), v
            yield from walk_key_values(v)
    elif isinstance(x, list):
        for v in x:
            yield from walk_key_values(v)


def extract_original_text(chapter_obj: Any) -> str:
    if isinstance(chapter_obj, str):
        return chapter_obj.strip()

    priority_keys = [
        "original_text",
        "normalized_original_text",
        "normalized_original_english_text",
        "normalized_original_english_chapter_text",
        "chapter_text",
        "baseline_text",
        "original_english_text",
        "text",
        "content",
    ]

    if isinstance(chapter_obj, dict):
        for key in priority_keys:
            value = chapter_obj.get(key)
            if isinstance(value, str) and len(value.strip()) > 200:
                return value.strip()

    candidates = []
    for k, v in walk_key_values(chapter_obj):
        kl = k.lower()
        if not isinstance(v, str):
            continue
        if len(v.strip()) < 200:
            continue
        if any(bad in kl for bad in [
            "reference", "figure", "table", "methodology",
            "conflict", "inconsisten", "scope_note",
        ]):
            continue
        if "text" in kl or "content" in kl or "chapter" in kl:
            candidates.append(v.strip())

    if candidates:
        return max(candidates, key=len)

    raise RuntimeError(
        "Could not identify original chapter narrative text in baseline object. "
        "Inspect the chapter keys and add its field name to extract_original_text()."
    )


def normalize_reference_item(item: Any, index: int) -> tuple[int, str] | None:
    if isinstance(item, str):
        txt = clean(item)
        if not txt:
            return None
        m = re.match(r"^\s*(\d+)[\.\)]\s*(.+)$", txt)
        if m:
            return int(m.group(1)), m.group(2).strip()
        return index, txt

    if isinstance(item, dict):
        num = (
            item.get("number")
            or item.get("reference_number")
            or item.get("ref_number")
            or item.get("id")
            or index
        )
        try:
            num = int(str(num).strip())
        except Exception:
            num = index

        for key in [
            "citation", "reference", "text", "full_citation",
            "formatted_reference", "reference_text",
        ]:
            if clean(item.get(key)):
                return num, clean(item.get(key))

        fields = [
            clean(item.get("authors")),
            clean(item.get("title")),
            clean(item.get("journal")),
            clean(item.get("year")),
            clean(item.get("volume")),
            clean(item.get("pages")),
        ]
        txt = ". ".join(x for x in fields if x)
        if txt:
            return num, txt

    return None


def extract_original_references(baseline_raw: Any) -> list[str]:
    candidates = []

    for obj in walk_objects(baseline_raw):
        if not isinstance(obj, list):
            continue
        if len(obj) < 58:
            continue

        normalized = []
        for i, item in enumerate(obj, 1):
            n = normalize_reference_item(item, i)
            if n:
                normalized.append(n)

        nums = {n for n, _ in normalized}
        if all(i in nums for i in range(1, 59)):
            candidates.append(normalized)

    if not candidates:
        raise RuntimeError(
            "Could not find a complete original ESMO reference list [1]-[58] "
            "inside esmo2015_baseline_by_chapter.json."
        )

    best = max(candidates, key=len)
    by_num = {}
    for n, text in best:
        if 1 <= n <= 58 and n not in by_num:
            by_num[n] = text

    if len(by_num) != 58:
        raise RuntimeError(
            f"Original reference extraction found only {len(by_num)}/58 references."
        )

    return [by_num[i] for i in range(1, 59)]


def paper_metadata_lookup() -> dict[str, dict[str, Any]]:
    rows = load_jsonl(FINAL_ASSIGNMENTS)
    grouped = defaultdict(list)

    for row in rows:
        pmid = clean(row.get("pmid"))
        if pmid:
            grouped[pmid].append(row)

    lookup = {}
    for pmid, items in grouped.items():
        richest = {}
        for field in [
            "pmid", "doi", "pmcid", "title", "authors", "journal",
            "publication_date", "publication_year", "publication_types",
            "evidence_labels",
        ]:
            richest[field] = max(
                (clean(x.get(field)) for x in items),
                key=len,
                default="",
            )
        lookup[pmid] = richest
    return lookup


def unique_in_order(values: Iterable[str]) -> list[str]:
    seen = set()
    out = []
    for x in values:
        x = clean(x)
        if x and x not in seen:
            seen.add(x)
            out.append(x)
    return out


def reference_appearance_order(results: list[dict[str, Any]]) -> list[str]:
    main_body = []
    appendix = []

    for row in sorted(results, key=lambda x: CHAPTER_ORDER.index(x["chapter_id"])):
        r = row["result"]
        for section in r.get("updated_sections") or []:
            for p in section.get("update_paragraphs") or []:
                main_body.extend(p.get("citation_pmids") or [])
            for p in section.get("current_clinical_practice") or []:
                main_body.extend(p.get("citation_pmids") or [])

        for p in r.get("chapter_current_clinical_practice_summary") or []:
            main_body.extend(p.get("citation_pmids") or [])

    for row in sorted(results, key=lambda x: CHAPTER_ORDER.index(x["chapter_id"])):
        r = row["result"]
        for topic in r.get("appendix_topics") or []:
            appendix.extend(topic.get("citation_pmids") or [])

    return unique_in_order(main_body + appendix)


def format_new_reference(meta: dict[str, Any]) -> str:
    authors = clean(meta.get("authors"))
    title = clean(meta.get("title"))
    journal = clean(meta.get("journal"))
    year = clean(meta.get("publication_year")) or clean(meta.get("publication_date"))
    doi = clean(meta.get("doi"))
    pmid = clean(meta.get("pmid"))

    parts = []
    if authors:
        parts.append(authors.rstrip("."))
    if title:
        parts.append(title.rstrip("."))
    journal_year = " ".join(x for x in [journal, year] if x).strip()
    if journal_year:
        parts.append(journal_year.rstrip("."))

    citation = ". ".join(parts).strip()
    if citation:
        citation += "."
    if doi:
        citation += f" doi:{doi}."
    if pmid:
        citation += f" PMID:{pmid}."
    return clean(citation)


def build_reference_registry(
    results: list[dict[str, Any]],
    baseline_raw: Any,
) -> dict[str, Any]:
    original_refs = extract_original_references(baseline_raw)
    metadata = paper_metadata_lookup()
    ordered_pmids = reference_appearance_order(results)

    new_refs = []
    pmid_to_number = {}

    next_num = 59
    for pmid in ordered_pmids:
        if pmid not in metadata:
            raise RuntimeError(
                f"Stage-B cited PMID {pmid} not found in final evidence assignments."
            )
        pmid_to_number[pmid] = next_num
        new_refs.append({
            "number": next_num,
            "pmid": pmid,
            "citation": format_new_reference(metadata[pmid]),
            "metadata": metadata[pmid],
        })
        next_num += 1

    registry = {
        "original_reference_numbers": [1, 58],
        "original_references": [
            {"number": i, "citation": original_refs[i - 1]}
            for i in range(1, 59)
        ],
        "new_reference_start": 59,
        "new_references": new_refs,
        "pmid_to_new_reference_number": pmid_to_number,
        "numbering_policy": (
            "New references begin at [59] and are numbered by first appearance "
            "in the Stage-B main narrative/current-practice sections, followed "
            "by appendix-only citations."
        ),
    }

    REFERENCE_REGISTRY.write_text(
        json.dumps(registry, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return registry


def citation_string(pmids: Iterable[str], registry: dict[str, Any]) -> str:
    mapping = registry["pmid_to_new_reference_number"]
    nums = []
    for p in unique_in_order(pmids):
        if p not in mapping:
            raise RuntimeError(f"No reference number registered for PMID {p}")
        nums.append(mapping[p])
    if not nums:
        return ""
    return " [" + ", ".join(str(n) for n in nums) + "]"


def ensure_docx_dependency():
    try:
        import docx  # noqa: F401
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for DOCX generation. Run this script with:\n"
            "  uv run --with python-docx python .\\src\\stageB_rewrite_and_docx.py --mode docx"
        ) from e


def set_cell_shading(cell, fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:{}".format(edge)
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn("w:{}".format(key)), str(edge_data[key]))


def add_cited_paragraph(doc, text: str, pmids: list[str], registry: dict[str, Any]):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = docx_shared_pt(6)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text.strip())
    r.font.name = "Times New Roman"
    r.font.size = docx_shared_pt(10.5)

    cites = citation_string(pmids, registry)
    if cites:
        r2 = p.add_run(cites)
        r2.font.name = "Times New Roman"
        r2.font.size = docx_shared_pt(10.5)
        r2.font.color.rgb = docx_rgb("8A1748")
    return p


def docx_shared_pt(value):
    from docx.shared import Pt
    return Pt(value)


def docx_rgb(hexstr: str):
    from docx.shared import RGBColor
    return RGBColor.from_string(hexstr)


def add_label_paragraph(doc, text: str, color: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = docx_shared_pt(8)
    p.paragraph_format.space_after = docx_shared_pt(4)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.name = "Arial"
    r.font.size = docx_shared_pt(9)
    r.font.color.rgb = docx_rgb(color)
    return p


def add_original_text(doc, text: str):
    paragraphs = [
        x.strip()
        for x in re.split(r"\n\s*\n|\r\n\s*\r\n", text)
        if x.strip()
    ]
    if len(paragraphs) == 1:
        # Some JSON baselines store chapter text with single newlines only.
        rough = [x.strip() for x in text.split("\n") if x.strip()]
        if len(rough) > 3:
            paragraphs = rough

    for block in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = docx_shared_pt(5)
        p.paragraph_format.line_spacing = 1.04
        r = p.add_run(block)
        r.font.name = "Times New Roman"
        r.font.size = docx_shared_pt(9.5)


def add_practice_box(
    doc,
    points: list[dict[str, Any]],
    registry: dict[str, Any],
    title: str = "Current clinical practice",
):
    if not points:
        return

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F5E8")
    set_cell_border(
        cell,
        left={"val": "single", "sz": "18", "color": "6F8121"},
        top={"val": "single", "sz": "4", "color": "D9DEC2"},
        bottom={"val": "single", "sz": "4", "color": "D9DEC2"},
        right={"val": "single", "sz": "4", "color": "D9DEC2"},
    )

    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = docx_shared_pt(10)
    r.font.color.rgb = docx_rgb("006633")

    for point in points:
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = docx_shared_pt(3)
        r = p.add_run(clean(point.get("text")))
        r.font.name = "Times New Roman"
        r.font.size = docx_shared_pt(9.5)

        cites = citation_string(point.get("citation_pmids") or [], registry)
        if cites:
            rc = p.add_run(cites)
            rc.font.name = "Times New Roman"
            rc.font.size = docx_shared_pt(9.5)
            rc.font.color.rgb = docx_rgb("8A1748")

        status = clean(point.get("practice_status"))
        if status:
            rs = p.add_run(f"  ({status.replace('_', ' ').title()})")
            rs.italic = True
            rs.font.name = "Arial"
            rs.font.size = docx_shared_pt(8)
            rs.font.color.rgb = docx_rgb("666666")

    doc.add_paragraph()


def configure_document(doc):
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = docx_shared_pt(10.5)

    for level, size in [(1, 16), (2, 13), (3, 11)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = docx_shared_pt(size)
        style.font.color.rgb = docx_rgb("006633")

    # Header
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("ESMO pancreatic cancer | Living evidence update")
    r.font.name = "Arial"
    r.font.size = docx_shared_pt(8)
    r.font.color.rgb = docx_rgb("777777")

    # Footer
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Evidence search through 31 August 2023 | Figures deferred to v2")
    r.font.name = "Arial"
    r.font.size = docx_shared_pt(8)
    r.font.color.rgb = docx_rgb("6F8121")


def add_title_page(doc):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = docx_shared_pt(60)

    r = p.add_run(
        "Cancer of the pancreas:\n"
        "ESMO 2015 Clinical Practice Guideline\n"
        "Living Evidence Update through August 2023"
    )
    r.bold = True
    r.font.name = "Arial"
    r.font.size = docx_shared_pt(21)
    r.font.color.rgb = docx_rgb("006633")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Evidence-integrated proof-of-concept | Text-first version 1")
    r.font.name = "Arial"
    r.font.size = docx_shared_pt(12)
    r.font.color.rgb = docx_rgb("8A1748")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = docx_shared_pt(28)
    r = p.add_run(
        "Document structure: original ESMO 2015 content followed by the "
        "2015-2023 evidence update and a concise current-clinical-practice summary."
    )
    r.font.name = "Times New Roman"
    r.font.size = docx_shared_pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Figures and treatment algorithms are intentionally deferred to a second run."
    )
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = docx_shared_pt(10)
    r.font.color.rgb = docx_rgb("666666")

    doc.add_page_break()


def add_summary_table(doc, results, registry):
    doc.add_heading("Summary of current clinical practice", level=1)
    p = doc.add_paragraph(
        "The following points summarize the evidence-integrated clinical practice "
        "interpretation based on literature through 31 August 2023."
    )
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = docx_shared_pt(10.5)

    for row in results:
        cid = row["chapter_id"]
        r = row["result"]
        points = r.get("chapter_current_clinical_practice_summary") or []
        if not points:
            continue
        doc.add_heading(
            f"{cid}  {CHAPTER_FALLBACK_TITLES[cid]}",
            level=2,
        )
        add_practice_box(doc, points, registry, title="Current clinical practice")


def add_references(doc, registry):
    doc.add_page_break()
    doc.add_heading("References", level=1)

    add_label_paragraph(doc, "Original ESMO 2015 references", "6F8121")
    for ref in registry["original_references"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = docx_shared_pt(12)
        p.paragraph_format.first_line_indent = docx_shared_pt(-12)
        p.paragraph_format.space_after = docx_shared_pt(2)
        r = p.add_run(f"[{ref['number']}] {ref['citation']}")
        r.font.name = "Times New Roman"
        r.font.size = docx_shared_pt(8.8)

    if registry["new_references"]:
        add_label_paragraph(doc, "Updated evidence 2015-2023", "8A1748")
        for ref in registry["new_references"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = docx_shared_pt(12)
            p.paragraph_format.first_line_indent = docx_shared_pt(-12)
            p.paragraph_format.space_after = docx_shared_pt(2)
            r = p.add_run(f"[{ref['number']}] {ref['citation']}")
            r.font.name = "Times New Roman"
            r.font.size = docx_shared_pt(8.8)


def add_appendix(doc, results, registry):
    appendix_rows = []
    for row in results:
        for topic in row["result"].get("appendix_topics") or []:
            appendix_rows.append((row["chapter_id"], topic))

    if not appendix_rows:
        return

    doc.add_page_break()
    doc.add_heading(
        "Appendix A - Emerging and non-recommendation-driving evidence",
        level=1,
    )

    p = doc.add_paragraph(
        "This appendix summarizes evidence retained for scientific context or "
        "future updates but not used independently to define current clinical practice."
    )
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = docx_shared_pt(10)

    current_chapter = None
    for cid, topic in appendix_rows:
        if cid != current_chapter:
            current_chapter = cid
            doc.add_heading(
                f"{cid}  {CHAPTER_FALLBACK_TITLES[cid]}",
                level=2,
            )

        doc.add_heading(clean(topic.get("title")), level=3)
        add_cited_paragraph(
            doc,
            clean(topic.get("summary")),
            topic.get("citation_pmids") or [],
            registry,
        )


def build_docx(appendix_mode: str) -> None:
    ensure_docx_dependency()
    from docx import Document

    m = load_json(MANIFEST)
    if m.get("status") != "READY_FOR_DOCX":
        raise RuntimeError(
            f"Stage B is not READY_FOR_DOCX: {m.get('status')}"
        )

    baseline_raw = load_json(BASELINE)
    baseline = baseline_by_chapter(baseline_raw)
    results = load_jsonl(RESULTS)
    if len(results) != 8:
        raise RuntimeError(f"Expected 8 chapter results, found {len(results)}")

    results.sort(key=lambda x: CHAPTER_ORDER.index(x["chapter_id"]))
    registry = build_reference_registry(results, baseline_raw)

    OUTPUT.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    # Methods note.
    doc.add_heading("Document convention", level=1)
    p = doc.add_paragraph(
        "For transparency, each chapter retains the original 2015 ESMO content "
        "with its original reference numbering [1-58]. The subsequent evidence-"
        "update section summarizes newly integrated evidence through 31 August "
        "2023. New references are numbered from [59] onward. The current-clinical-"
        "practice box provides the practical interpretation of the integrated "
        "evidence. Figures are intentionally omitted from this version."
    )
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = docx_shared_pt(10.5)

    for row in results:
        cid = row["chapter_id"]
        r = row["result"]

        doc.add_page_break()
        doc.add_heading(
            f"{cid}  {CHAPTER_FALLBACK_TITLES[cid]}",
            level=1,
        )

        add_label_paragraph(doc, "Original ESMO 2015", "6F8121")
        original_text = extract_original_text(baseline[cid])
        add_original_text(doc, original_text)

        add_label_paragraph(
            doc,
            "Evidence update 2015-31 August 2023",
            "8A1748",
        )

        p = doc.add_paragraph()
        rr = p.add_run(clean(r.get("overall_update_summary")))
        rr.bold = True
        rr.font.name = "Times New Roman"
        rr.font.size = docx_shared_pt(10.5)

        for section in r.get("updated_sections") or []:
            doc.add_heading(clean(section.get("section_heading")), level=2)

            anchor = clean(section.get("anchor_in_original"))
            signal = clean(section.get("change_signal"))
            if anchor or signal:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = docx_shared_pt(4)
                ra = p.add_run(
                    "2015 anchor: " + (anchor or "chapter-level")
                    + " | Evidence signal: " + signal.replace("_", " ").title()
                )
                ra.italic = True
                ra.font.name = "Arial"
                ra.font.size = docx_shared_pt(8.5)
                ra.font.color.rgb = docx_rgb("666666")

            for para in section.get("update_paragraphs") or []:
                add_cited_paragraph(
                    doc,
                    clean(para.get("text")),
                    para.get("citation_pmids") or [],
                    registry,
                )

            interpretation = clean(section.get("interpretation_of_2015_position"))
            if interpretation:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = docx_shared_pt(6)
                ri = p.add_run("Interpretation of the 2015 position: ")
                ri.bold = True
                ri.font.name = "Arial"
                ri.font.size = docx_shared_pt(9)
                rt = p.add_run(interpretation)
                rt.font.name = "Times New Roman"
                rt.font.size = docx_shared_pt(9.5)

            add_practice_box(
                doc,
                section.get("current_clinical_practice") or [],
                registry,
                title="Current clinical practice",
            )

        chapter_points = r.get("chapter_current_clinical_practice_summary") or []
        if chapter_points:
            add_practice_box(
                doc,
                chapter_points,
                registry,
                title="Chapter summary - current clinical practice",
            )

        limitations = clean(r.get("chapter_limitations"))
        if limitations:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = docx_shared_pt(6)
            rr = p.add_run("Evidence limitations: ")
            rr.bold = True
            rr.font.name = "Arial"
            rr.font.size = docx_shared_pt(9)
            rt = p.add_run(limitations)
            rt.font.name = "Times New Roman"
            rt.font.size = docx_shared_pt(9)

    doc.add_page_break()
    add_summary_table(doc, results, registry)

    if appendix_mode in {"main", "both"}:
        add_appendix(doc, results, registry)

    add_references(doc, registry)

    doc.save(DOCX_PATH)

    if appendix_mode in {"separate", "both"}:
        app = Document()
        configure_document(app)
        p = app.add_paragraph()
        r = p.add_run(
            "Appendix - Emerging and non-recommendation-driving evidence"
        )
        r.bold = True
        r.font.name = "Arial"
        r.font.size = docx_shared_pt(18)
        r.font.color.rgb = docx_rgb("006633")
        add_appendix(app, results, registry)
        add_references(app, registry)
        app.save(APPENDIX_DOCX_PATH)

    print("\nDOCX generation completed.")
    print(f"  main DOCX:       {DOCX_PATH}")
    print(f"  appendix mode:   {appendix_mode}")
    if appendix_mode in {"separate", "both"}:
        print(f"  appendix DOCX:   {APPENDIX_DOCX_PATH}")
    print(f"  ref registry:    {REFERENCE_REGISTRY}")
    print()
    print(
        "Figures/flowcharts were intentionally not generated in this v1 document."
    )


def synchronous_test(client: Client):
    with BATCH_INPUT.open("r", encoding="utf-8") as f:
        req = json.loads(f.readline())
    print(f"Testing: {req['custom_id']}")
    r = client.request(
        "POST",
        "/chat/completions",
        headers={"Content-Type": "application/json"},
        data=json.dumps(req["body"]),
        timeout=1800,
    )
    print(f"HTTP {r.status_code}")
    data = r.json()
    choice = (data.get("choices") or [{}])[0]
    print(f"finish_reason: {choice.get('finish_reason')}")
    print(
        "content length:",
        len(((choice.get("message") or {}).get("content") or "")),
    )
    print(json.dumps(data, ensure_ascii=False, indent=2)[:16000])


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--mode",
        choices=[
            "prepare",
            "test",
            "submit",
            "watch",
            "merge",
            "docx",
            "all",
        ],
        default="prepare",
    )
    parser.add_argument("--model", default="gpt-5.6-sol")
    parser.add_argument(
        "--reasoning-effort",
        choices=["none", "low", "medium", "high", "xhigh", "max"],
        default="high",
    )
    parser.add_argument("--poll-seconds", type=int, default=300)
    parser.add_argument("--retry-wait", type=int, default=120)
    parser.add_argument(
        "--appendix-mode",
        choices=["main", "separate", "both", "none"],
        default="main",
    )
    parser.add_argument("--reset-state", action="store_true")
    args = parser.parse_args()

    if args.reset_state and STATE.exists():
        STATE.unlink()
        print(f"Deleted: {STATE}")

    if args.mode in {"prepare", "all"}:
        prepare(args.model, args.reasoning_effort)
        if args.mode == "prepare":
            return 0

    if args.mode == "merge":
        merge()
        return 0

    if args.mode == "docx":
        build_docx(args.appendix_mode)
        return 0

    api_key = os.environ.get("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set.")
    client = Client(api_key, args.retry_wait)

    if args.mode == "test":
        synchronous_test(client)
        return 0

    if args.mode in {"submit", "all"}:
        submit(client, args.model, args.reasoning_effort)
        if args.mode == "submit":
            return 0

    if args.mode in {"watch", "all"}:
        batch = watch(client, args.poll_seconds)
        print(f"Batch terminal status: {batch.get('status')}")
        if batch.get("status") != "completed":
            return 2

        merged = merge()
        if merged.get("status") != "READY_FOR_DOCX":
            return 3

        build_docx(args.appendix_mode)
        return 0

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
